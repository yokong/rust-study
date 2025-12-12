import sys
import os
import json
import subprocess

# SCAN_DIR = os.path.dirname(__file__)
SCAN_DIR = '/Volumes/安全硬盘小/APP安全检测报告'
KEYWORDS = ['检测', 'HYA', 'secidea']
CASE_INSENSITIVE = True

def parse_keywords(s):
    try:
        v = json.loads(s)
        if isinstance(v, list):
            return [str(x) for x in v if str(x)]
    except Exception:
        pass
    return [x for x in [y.strip() for y in s.split(',')] if x]

def find_files(root):
    res = []
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            lf = fn.lower()
            if lf.endswith('.docx') or lf.endswith('.doc'):
                res.append(os.path.join(dirpath, fn))
    return res

def search_indices(text, kw, ci):
    t = text if not ci else text.lower()
    k = kw if not ci else kw.lower()
    i = 0
    out = []
    while True:
        i = t.find(k, i)
        if i == -1:
            break
        out.append(i)
        i += len(k) if len(k) > 0 else 1
    return out

def process_docx(path, keywords, ci):
    try:
        import docx
    except Exception:
        return process_via_textutil(path, keywords, ci)
    res = []
    try:
        d = docx.Document(path)
    except Exception:
        return res
    def scan_container(container, area_prefix):
        for idx, p in enumerate(getattr(container, 'paragraphs', []) or []):
            txt = p.text or ''
            for kw in keywords:
                for pos in search_indices(txt, kw, ci):
                    res.append({'file': path, 'area': area_prefix, 'detail': f'paragraph {idx}', 'keyword': kw, 'index': pos, 'snippet': txt[max(0,pos-20):pos+20]})
        for tr_idx, tbl in enumerate(getattr(container, 'tables', []) or []):
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        txt = p.text or ''
                        for kw in keywords:
                            for pos in search_indices(txt, kw, ci):
                                res.append({'file': path, 'area': area_prefix, 'detail': f'table {tr_idx} row {r_idx} cell {c_idx} para {p_idx}', 'keyword': kw, 'index': pos, 'snippet': txt[max(0,pos-20):pos+20]})

    scan_container(d, 'body')
    for s_idx, sec in enumerate(d.sections):
        for hdr in [getattr(sec, 'header', None), getattr(sec, 'first_page_header', None), getattr(sec, 'even_page_header', None)]:
            if hdr is not None:
                scan_container(hdr, 'header')
        for ftr in [getattr(sec, 'footer', None), getattr(sec, 'first_page_footer', None), getattr(sec, 'even_page_footer', None)]:
            if ftr is not None:
                scan_container(ftr, 'footer')
    return res

def process_docx_properties(path, keywords, ci):
    try:
        import docx
    except Exception:
        return []
    res = []
    try:
        d = docx.Document(path)
    except Exception:
        return res
    cp = d.core_properties
    fields = {
        'title': getattr(cp, 'title', None),
        'subject': getattr(cp, 'subject', None),
        'author': getattr(cp, 'author', None),
        'category': getattr(cp, 'category', None),
        'comments': getattr(cp, 'comments', None),
        'keywords': getattr(cp, 'keywords', None),
        'last_modified_by': getattr(cp, 'last_modified_by', None),
    }
    for name, val in fields.items():
        if not val:
            continue
        txt = str(val)
        for kw in keywords:
            for pos in search_indices(txt, kw, ci):
                res.append({'file': path, 'area': 'properties', 'detail': f'prop {name}', 'keyword': kw, 'index': pos, 'snippet': txt[max(0,pos-20):pos+20]})
    return res

def process_via_textutil(path, keywords, ci):
    res = []
    try:
        p = subprocess.run(['textutil', '-convert', 'txt', '-stdout', path], capture_output=True, check=True)
        txt = p.stdout.decode('utf-8', errors='replace')
    except Exception:
        return res
    for kw in keywords:
        for pos in search_indices(txt, kw, ci):
            res.append({'file': path, 'area': 'body', 'detail': 'textutil', 'keyword': kw, 'index': pos, 'snippet': txt[max(0,pos-20):pos+20]})
    return res

def process_doc_properties_mdls(path, keywords, ci):
    res = []
    try:
        p = subprocess.run(['mdls', '-name', 'kMDItemTitle', '-name', 'kMDItemKeywords', '-name', 'kMDItemAuthors', '-name', 'kMDItemDescription', path], capture_output=True, check=True)
        out = p.stdout.decode('utf-8', errors='replace')
    except Exception:
        return res
    collected = []
    for line in out.splitlines():
        if ':' not in line:
            continue
        k, v = line.split(':', 1)
        v = v.strip()
        if v and v != '(null)':
            collected.append((k.strip(), v))
    for name, txt in collected:
        for kw in keywords:
            for pos in search_indices(txt, kw, ci):
                res.append({'file': path, 'area': 'properties', 'detail': f'mdls {name}', 'keyword': kw, 'index': pos, 'snippet': txt[max(0,pos-20):pos+20]})
    return res

def write_result(root, items):
    out_path = os.path.join(root, 'reslut.txt')
    by_file = {}
    for it in items:
        fp = it['file']
        by_file.setdefault(fp, []).append(it)
    def area_cn(a):
        if a == 'body':
            return '正文'
        if a == 'header':
            return '页眉'
        if a == 'footer':
            return '页脚'
        if a == 'properties':
            return '文档属性'
        return a
    with open(out_path, 'w', encoding='utf-8') as f:
        for fp, hits in by_file.items():
            kws = sorted(set(h['keyword'] for h in hits))
            cats_cn = sorted(set(area_cn(h['area']) for h in hits))
            f.write("=================\n")
            f.write(f"文件路径：{fp}\n")
            f.write(f"命中关键词：{(' | '.join(kws) if kws else '无')}\n")
            f.write(f"命中类别：{(' | '.join(cats_cn) if cats_cn else '无')}\n")
            f.write("命中位置：\n")
            for h in hits:
                f.write(f"- {area_cn(h['area'])}：{h['detail']} @ {h['index']}\n")
            f.write("\n")
    return out_path

def main():
    ci = CASE_INSENSITIVE
    keywords = KEYWORDS
    files = find_files(SCAN_DIR)
    total = len(files)
    print(f"开始扫描: {SCAN_DIR}，文件数: {total}")
    all_items = []
    for i, fp in enumerate(files, 1):
        print(f"[{i}/{total}] 扫描 {fp}")
        if fp.lower().endswith('.docx'):
            found = []
            found.extend(process_docx(fp, keywords, ci))
            found.extend(process_docx_properties(fp, keywords, ci))
            all_items.extend(found)
        else:
            found = []
            found.extend(process_via_textutil(fp, keywords, ci))
            found.extend(process_doc_properties_mdls(fp, keywords, ci))
            all_items.extend(found)
        body_hits = sum(1 for x in found if x['area'] in ('body'))
        header_footer_hits = sum(1 for x in found if x['area'] in ('header','footer'))
        properties_hits = sum(1 for x in found if x['area'] in ('properties'))
        hit_keywords = sorted(set(x['keyword'] for x in found))
        print(f"[{i}/{total}] 完成")
        print(f" 页眉页脚：命中 {header_footer_hits}")
        print(f" 正文：命中 {body_hits}")
        print(f" 文档属性：命中 {properties_hits}")
        print(f" 命中关键词：{', '.join(hit_keywords) if hit_keywords else '无'}")
    out_path = write_result(os.path.dirname(__file__), all_items)
    print(f"结果输出: {out_path}")

if __name__ == '__main__':
    main()
